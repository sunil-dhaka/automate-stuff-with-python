import docx

doc=docx.Document('files/demo.docx')

print('paragraphs'.center(50,'='))
for para in doc.paragraphs:
    print(para.text)
print('runs'.center(50,'='))
for para in doc.paragraphs:
    for run in para.runs:
        print(run.text)
