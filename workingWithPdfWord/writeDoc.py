import docx,sys

from pandas import wide_to_long

def createDocx(file):

    with open(file,'r') as file:
        guests=[line.strip() for line in file.readlines()]

    new_doc=docx.Document()
    for guest in guests:
        new_doc.add_heading(f'Dear {guest},',level=1).bold=True
        body=new_doc.add_paragraph('I hope that you are doing better than me. On ','BodyText')
        date=body.add_run('7th April, 2022')
        date.italic=True
        date.bold=True
        body.add_run(', there is a special')
        underline_text=body.add_run(' environment friendly')
        underline_text.underline=True
        underline_text.emboss=True
        body.add_run(' party at ')
        location=body.add_run('AP0021, Kalmfins Block, Jamupa City')
        location.shadow=True
        location.outline=True
        new_doc.add_paragraph('Please bring somethign to eat 😀. Some ideas:')
        new_doc.add_paragraph('Bread and Wine',style='List Bullet')
        new_doc.add_paragraph('Cake',style='List Bullet')
        new_doc.add_paragraph('Emrosia',style='List Bullet')
        new_doc.add_paragraph('Somras',style='List Bullet')
        new_doc.add_paragraph('Nectar',style='List Bullet')
        new_doc.add_heading('Please see below invitation image.',level=3)
        new_doc.add_picture('files/example_image.png',width=docx.shared.Inches(4),height=docx.shared.Cm(12))
        new_doc.add_paragraph('Tit for tats.',style='Intense Quote')
        new_doc.add_page_break()
    new_doc.save('invitation.docx')

if __name__=="__main__":
    if len(sys.argv)>1:
        file=sys.argv[1] 
    else:
        print('Give <file>.')
        sys.exit()
    
    createDocx(file)